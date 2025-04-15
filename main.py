import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
from PIL import Image
import pandas as pd
import json
import csv
from pathlib import Path
import docx
from docx import Document
import PyPDF2
import markdown
import html2text
import yaml
import striprtf.striprtf as striprtf


class FileConverterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("File Format Converter")
        self.root.geometry("680x300")
        self.root.configure(bg="#f0f0f0")
        
        # Supported conversions
        self.supported_conversions = {
            "Image": {
                "PNG": [".jpg", ".jpeg", ".bmp", ".gif", ".webp"],
                "JPG": [".png", ".bmp", ".gif", ".webp"],
                "BMP": [".png", ".jpg", ".jpeg", ".gif", ".webp"],
                "GIF": [".png", ".jpg", ".jpeg", ".bmp", ".webp"],
                "WEBP": [".png", ".jpg", ".jpeg", ".bmp", ".gif"]
            },
            "Data": {
                "CSV": [".xlsx", ".json", ".yaml", ".txt"],
                "JSON": [".csv", ".xlsx", ".yaml", ".txt"],
                "XLSX": [".csv", ".json", ".yaml", ".txt"],
                "YAML": [".csv", ".json", ".xlsx", ".txt"],
                "TXT": [".csv", ".json", ".xlsx", ".yaml"]
            },
            "Document": {
                "DOCX": [".txt", ".pdf", ".md", ".html", ".rtf"],
                "PDF": [".txt", ".docx", ".md", ".html", ".rtf"],
                "TXT": [".docx", ".pdf", ".md", ".html", ".rtf"],
                "MD": [".txt", ".docx", ".pdf", ".html", ".rtf"],
                "HTML": [".txt", ".docx", ".pdf", ".md", ".rtf"],
                "RTF": [".txt", ".docx", ".pdf", ".md", ".html"]
            }
        }
        
        self.setup_ui()
        
    def setup_ui(self):
        # Main frame
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Title
        title_label = ttk.Label(main_frame, text="File Format Converter", font=("Helvetica", 16, "bold"))
        title_label.grid(row=0, column=0, columnspan=2, pady=20)
        
        # File selection
        ttk.Label(main_frame, text="Select File:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.file_path = tk.StringVar()
        file_entry = ttk.Entry(main_frame, textvariable=self.file_path, width=50)
        file_entry.grid(row=1, column=1, padx=5, pady=5)
        browse_btn = ttk.Button(main_frame, text="Browse", command=self.browse_file)
        browse_btn.grid(row=1, column=2, padx=5, pady=5)
        
        # Format selection
        ttk.Label(main_frame, text="Convert to:").grid(row=2, column=0, sticky=tk.W, pady=5)
        self.target_format = tk.StringVar()
        self.format_combo = ttk.Combobox(main_frame, textvariable=self.target_format, state="readonly")
        self.format_combo.grid(row=2, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Convert button
        convert_btn = ttk.Button(main_frame, text="Convert", command=self.convert_file)
        convert_btn.grid(row=3, column=0, columnspan=3, pady=20)
        
        # Progress bar
        self.progress = ttk.Progressbar(main_frame, length=400, mode='determinate')
        self.progress.grid(row=4, column=0, columnspan=3, pady=10)
        
        # Status label
        self.status_var = tk.StringVar()
        self.status_var.set("Ready")
        status_label = ttk.Label(main_frame, textvariable=self.status_var)
        status_label.grid(row=5, column=0, columnspan=3, pady=5)
        
    def browse_file(self):
        file_path = filedialog.askopenfilename()
        if file_path:
            self.file_path.set(file_path)
            self.update_target_formats(file_path)
            
    def update_target_formats(self, file_path):
        ext = os.path.splitext(file_path)[1].lower()
        target_formats = []
        
        # Find supported conversions for the file type
        for category, formats in self.supported_conversions.items():
            for source_format, targets in formats.items():
                if f".{source_format.lower()}" == ext:
                    target_formats = [fmt.strip('.') for fmt in targets]
                    break
                    
        if target_formats:
            self.format_combo['values'] = target_formats
            if target_formats:
                self.format_combo.set(target_formats[0])
        else:
            self.format_combo['values'] = []
            self.format_combo.set('')
            messagebox.showwarning("Unsupported Format", "This file format is not supported for conversion.")
            
    def convert_file(self):
        source_path = self.file_path.get()
        if not source_path:
            messagebox.showerror("Error", "Please select a file to convert")
            return
            
        target_format = self.target_format.get()
        if not target_format:
            messagebox.showerror("Error", "Please select a target format")
            return
            
        try:
            self.progress['value'] = 0
            self.status_var.set("Converting...")
            
            source_ext = os.path.splitext(source_path)[1].lower()
            target_path = os.path.splitext(source_path)[0] + f".{target_format.lower()}"
            
            # Image conversion
            if source_ext in ['.png', '.jpg', '.jpeg', '.bmp', '.gif', '.webp']:
                img = Image.open(source_path)
                # Convert RGBA to RGB if saving as JPEG
                if target_format.lower() in ['jpg', 'jpeg'] and img.mode == 'RGBA':
                    # Create a white background
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    # Paste the image onto the background using alpha channel as mask
                    background.paste(img, mask=img.split()[3])
                    img = background
                img.save(target_path)
                
            # Data conversion
            elif source_ext == '.csv':
                df = pd.read_csv(source_path)
                if target_format == 'JSON':
                    df.to_json(target_path, orient='records')
                elif target_format == 'XLSX':
                    df.to_excel(target_path, index=False)
                elif target_format == 'YAML':
                    df.to_yaml(target_path)
                elif target_format == 'TXT':
                    df.to_csv(target_path, sep='\t', index=False)
                    
            elif source_ext == '.json':
                df = pd.read_json(source_path)
                if target_format == 'CSV':
                    df.to_csv(target_path, index=False)
                elif target_format == 'XLSX':
                    df.to_excel(target_path, index=False)
                elif target_format == 'YAML':
                    df.to_yaml(target_path)
                elif target_format == 'TXT':
                    df.to_csv(target_path, sep='\t', index=False)
                    
            elif source_ext == '.xlsx':
                df = pd.read_excel(source_path)
                if target_format == 'CSV':
                    df.to_csv(target_path, index=False)
                elif target_format == 'JSON':
                    df.to_json(target_path, orient='records')
                elif target_format == 'YAML':
                    df.to_yaml(target_path)
                elif target_format == 'TXT':
                    df.to_csv(target_path, sep='\t', index=False)
                    
            elif source_ext == '.yaml':
                df = pd.read_yaml(source_path)
                if target_format == 'CSV':
                    df.to_csv(target_path, index=False)
                elif target_format == 'JSON':
                    df.to_json(target_path, orient='records')
                elif target_format == 'XLSX':
                    df.to_excel(target_path, index=False)
                elif target_format == 'TXT':
                    df.to_csv(target_path, sep='\t', index=False)
                    
            elif source_ext == '.txt':
                with open(source_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                if target_format == 'CSV':
                    # Try to parse as CSV and convert
                    try:
                        df = pd.read_csv(source_path, sep=None, engine='python')
                        df.to_csv(target_path, index=False)
                    except:
                        messagebox.showerror("Error", "Could not convert text file to CSV format")
                        return
                elif target_format == 'JSON':
                    # Try to parse as JSON and convert
                    try:
                        df = pd.read_json(source_path)
                        df.to_json(target_path, orient='records')
                    except:
                        messagebox.showerror("Error", "Could not convert text file to JSON format")
                        return
                elif target_format == 'XLSX':
                    # Try to parse as table and convert
                    try:
                        df = pd.read_csv(source_path, sep=None, engine='python')
                        df.to_excel(target_path, index=False)
                    except:
                        messagebox.showerror("Error", "Could not convert text file to Excel format")
                        return
                elif target_format == 'YAML':
                    # Try to parse as YAML and convert
                    try:
                        df = pd.read_yaml(source_path)
                        df.to_yaml(target_path)
                    except:
                        messagebox.showerror("Error", "Could not convert text file to YAML format")
                        return
                elif target_format == 'RTF':
                    # Convert text to RTF
                    rtf_content = "{\\rtf1\\ansi\\deff0\n"
                    for line in content.split('\n'):
                        rtf_content += line + "\\par\n"
                    rtf_content += "}"
                    with open(target_path, 'w', encoding='utf-8') as f:
                        f.write(rtf_content)
                        
            # Document conversion
            elif source_ext == '.docx':
                doc = Document(source_path)
                if target_format == 'TXT':
                    with open(target_path, 'w', encoding='utf-8') as f:
                        for para in doc.paragraphs:
                            f.write(para.text + '\n')
                elif target_format == 'PDF':
                    # Convert to PDF using docx2pdf
                    from docx2pdf import convert
                    convert(source_path, target_path)
                elif target_format == 'MD':
                    with open(target_path, 'w', encoding='utf-8') as f:
                        for para in doc.paragraphs:
                            f.write(para.text + '\n\n')
                elif target_format == 'HTML':
                    html_content = "<html><body>"
                    for para in doc.paragraphs:
                        html_content += f"<p>{para.text}</p>"
                    html_content += "</body></html>"
                    with open(target_path, 'w', encoding='utf-8') as f:
                        f.write(html_content)
                elif target_format == 'RTF':
                    # Convert DOCX to RTF using python-docx2rtf
                    from docx2rtf import convert
                    convert(source_path, target_path)
                    
            elif source_ext == '.pdf':
                if target_format == 'TXT':
                    with open(source_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text()
                        with open(target_path, 'w', encoding='utf-8') as f:
                            f.write(text)
                elif target_format == 'DOCX':
                    # Convert PDF to DOCX using pdf2docx
                    from pdf2docx import Converter
                    cv = Converter(source_path)
                    cv.convert(target_path)
                    cv.close()
                elif target_format == 'MD':
                    with open(source_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text()
                        with open(target_path, 'w', encoding='utf-8') as f:
                            f.write(text)
                elif target_format == 'HTML':
                    with open(source_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text()
                        html_content = f"<html><body><pre>{text}</pre></body></html>"
                        with open(target_path, 'w', encoding='utf-8') as f:
                            f.write(html_content)
                elif target_format == 'RTF':
                    # Convert PDF to RTF via DOCX
                    from pdf2docx import Converter
                    temp_docx = source_path + ".docx"
                    cv = Converter(source_path)
                    cv.convert(temp_docx)
                    cv.close()
                    from docx2rtf import convert
                    convert(temp_docx, target_path)
                    os.remove(temp_docx)
                    
            elif source_ext == '.md':
                with open(source_path, 'r', encoding='utf-8') as f:
                    md_content = f.read()
                if target_format == 'TXT':
                    with open(target_path, 'w', encoding='utf-8') as f:
                        f.write(md_content)
                elif target_format == 'DOCX':
                    doc = Document()
                    doc.add_paragraph(md_content)
                    doc.save(target_path)
                elif target_format == 'PDF':
                    # Convert markdown to HTML first
                    html = markdown.markdown(md_content)
                    # Then convert HTML to PDF using weasyprint
                    from weasyprint import HTML
                    HTML(string=html).write_pdf(target_path)
                elif target_format == 'HTML':
                    html = markdown.markdown(md_content)
                    with open(target_path, 'w', encoding='utf-8') as f:
                        f.write(html)
                elif target_format == 'RTF':
                    # Convert markdown to RTF via DOCX
                    doc = Document()
                    doc.add_paragraph(md_content)
                    temp_docx = source_path + ".docx"
                    doc.save(temp_docx)
                    from docx2rtf import convert
                    convert(temp_docx, target_path)
                    os.remove(temp_docx)
                        
            elif source_ext == '.html':
                with open(source_path, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                if target_format == 'TXT':
                    h = html2text.HTML2Text()
                    h.ignore_links = True
                    text = h.handle(html_content)
                    with open(target_path, 'w', encoding='utf-8') as f:
                        f.write(text)
                elif target_format == 'DOCX':
                    doc = Document()
                    h = html2text.HTML2Text()
                    h.ignore_links = True
                    text = h.handle(html_content)
                    doc.add_paragraph(text)
                    doc.save(target_path)
                elif target_format == 'PDF':
                    from weasyprint import HTML
                    HTML(string=html_content).write_pdf(target_path)
                elif target_format == 'MD':
                    h = html2text.HTML2Text()
                    h.ignore_links = True
                    text = h.handle(html_content)
                    with open(target_path, 'w', encoding='utf-8') as f:
                        f.write(text)
                elif target_format == 'RTF':
                    # Convert HTML to RTF via DOCX
                    doc = Document()
                    h = html2text.HTML2Text()
                    h.ignore_links = True
                    text = h.handle(html_content)
                    doc.add_paragraph(text)
                    temp_docx = source_path + ".docx"
                    doc.save(temp_docx)
                    from docx2rtf import convert
                    convert(temp_docx, target_path)
                    os.remove(temp_docx)
                    
            elif source_ext == '.rtf':
                # Read RTF content
                with open(source_path, 'r', encoding='utf-8') as f:
                    rtf_content = f.read()
                # Strip RTF formatting
                text = striprtf.rtf_to_text(rtf_content)
                
                if target_format == 'TXT':
                    with open(target_path, 'w', encoding='utf-8') as f:
                        f.write(text)
                elif target_format == 'DOCX':
                    doc = Document()
                    doc.add_paragraph(text)
                    doc.save(target_path)
                elif target_format == 'PDF':
                    # Convert RTF to PDF via DOCX
                    doc = Document()
                    doc.add_paragraph(text)
                    temp_docx = source_path + ".docx"
                    doc.save(temp_docx)
                    from docx2pdf import convert
                    convert(temp_docx, target_path)
                    os.remove(temp_docx)
                elif target_format == 'MD':
                    with open(target_path, 'w', encoding='utf-8') as f:
                        f.write(text)
                elif target_format == 'HTML':
                    html_content = "<html><body>"
                    for line in text.split('\n'):
                        html_content += f"<p>{line}</p>"
                    html_content += "</body></html>"
                    with open(target_path, 'w', encoding='utf-8') as f:
                        f.write(html_content)
                    
            self.progress['value'] = 100
            self.status_var.set("Conversion completed!")
            messagebox.showinfo("Success", f"File converted successfully!\nSaved as: {target_path}")
            
        except Exception as e:
            self.status_var.set("Error during conversion")
            messagebox.showerror("Error", f"An error occurred during conversion: {str(e)}")
            
if __name__ == "__main__":
    root = tk.Tk()
    app = FileConverterApp(root)
    root.mainloop()
